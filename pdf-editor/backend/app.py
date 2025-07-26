import os
import json
import uuid
from pathlib import Path
from flask import Flask, request, send_from_directory, jsonify, send_file
from flask_cors import CORS
from werkzeug.utils import secure_filename
from PyPDF2 import PdfReader, PdfWriter
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches
import io
from PIL import Image, ImageDraw, ImageFont
import base64

UPLOAD_FOLDER = 'uploads'
PROCESSED_FOLDER = 'processed'
TEMP_FOLDER = 'temp'
ALLOWED_EXTENSIONS = {'pdf', 'docx', 'doc'}

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(PROCESSED_FOLDER, exist_ok=True)
os.makedirs(TEMP_FOLDER, exist_ok=True)

app = Flask(__name__)
CORS(app)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['PROCESSED_FOLDER'] = PROCESSED_FOLDER
app.config['TEMP_FOLDER'] = TEMP_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 1024 * 1024 * 1024  # 1GB

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/upload', methods=['POST'])
def upload_file():
    """Upload PDF or Word document"""
    if 'file' not in request.files:
        return jsonify({'error': 'No file part'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'error': 'No selected file'}), 400
    
    if file and allowed_file(file.filename):
        filename = secure_filename(file.filename)
        unique_name = f"{uuid.uuid4()}_{filename}"
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], unique_name)
        file.save(file_path)
        
        # Get file info
        file_info = {
            'filename': unique_name,
            'original_name': filename,
            'size': os.path.getsize(file_path),
            'type': filename.rsplit('.', 1)[1].lower()
        }
        
        # If it's a PDF, get page count
        if file_info['type'] == 'pdf':
            try:
                with open(file_path, 'rb') as f:
                    reader = PdfReader(f)
                    file_info['page_count'] = len(reader.pages)
            except Exception as e:
                file_info['page_count'] = 0
        
        return jsonify(file_info), 200
    
    return jsonify({'error': 'Invalid file type'}), 400

@app.route('/files', methods=['GET'])
def list_files():
    """List all uploaded files"""
    files = []
    for filename in os.listdir(app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        if os.path.isfile(file_path):
            file_info = {
                'filename': filename,
                'original_name': filename.split('_', 1)[1] if '_' in filename else filename,
                'size': os.path.getsize(file_path),
                'type': filename.rsplit('.', 1)[1].lower() if '.' in filename else 'unknown'
            }
            files.append(file_info)
    
    return jsonify(files), 200

@app.route('/download/<filename>', methods=['GET'])
def download_file(filename):
    """Download a file"""
    folder = app.config['PROCESSED_FOLDER'] if request.args.get('processed') else app.config['UPLOAD_FOLDER']
    return send_from_directory(folder, filename, as_attachment=True)

@app.route('/preview/<filename>', methods=['GET'])
def preview_file(filename):
    """Get PDF preview (first page as image)"""
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    
    if not os.path.exists(file_path) or not filename.lower().endswith('.pdf'):
        return jsonify({'error': 'File not found or not a PDF'}), 404
    
    try:
        # Convert first page to image
        doc = fitz.open(file_path)
        page = doc[0]
        pix = page.get_pixmap(matrix=fitz.Matrix(1.5, 1.5))
        img_data = pix.tobytes("png")
        
        # Convert to base64
        img_base64 = base64.b64encode(img_data).decode()
        
        doc.close()
        
        return jsonify({
            'preview': f"data:image/png;base64,{img_base64}",
            'page_count': len(doc)
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/edit', methods=['POST'])
def edit_pdf():
    """Edit PDF with granular details"""
    data = request.json
    filename = data.get('filename')
    edits = data.get('edits', [])  # List of edit operations
    
    if not filename:
        return jsonify({'error': 'Missing filename'}), 400
    
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(input_path):
        return jsonify({'error': 'File not found'}), 404
    
    try:
        # Open PDF with PyMuPDF for editing
        doc = fitz.open(input_path)
        
        for edit in edits:
            page_num = edit.get('page', 1) - 1  # Convert to 0-based index
            edit_type = edit.get('type')
            
            if page_num >= len(doc):
                continue
            
            page = doc[page_num]
            
            if edit_type == 'text':
                # Add text
                text = edit.get('text', '')
                x, y = edit.get('x', 100), edit.get('y', 100)
                font_size = edit.get('font_size', 12)
                color = edit.get('color', (0, 0, 0))
                
                page.insert_text((x, y), text, fontsize=font_size, color=color)
            
            elif edit_type == 'image':
                # Add image
                image_data = edit.get('image_data', '')
                x, y = edit.get('x', 100), edit.get('y', 100)
                width = edit.get('width', 100)
                height = edit.get('height', 100)
                
                # Decode base64 image
                if image_data.startswith('data:image'):
                    image_data = image_data.split(',')[1]
                
                img_bytes = base64.b64decode(image_data)
                img_rect = fitz.Rect(x, y, x + width, y + height)
                page.insert_image(img_rect, stream=img_bytes)
            
            elif edit_type == 'rectangle':
                # Add rectangle
                x, y = edit.get('x', 100), edit.get('y', 100)
                width, height = edit.get('width', 100), edit.get('height', 100)
                color = edit.get('color', (0, 0, 0))
                fill = edit.get('fill', None)
                
                rect = fitz.Rect(x, y, x + width, y + height)
                page.draw_rect(rect, color=color, fill=fill)
            
            elif edit_type == 'line':
                # Add line
                x1, y1 = edit.get('x1', 100), edit.get('y1', 100)
                x2, y2 = edit.get('x2', 200), edit.get('y2', 200)
                color = edit.get('color', (0, 0, 0))
                width = edit.get('width', 1)
                
                page.draw_line((x1, y1), (x2, y2), color=color, width=width)
        
        # Save edited PDF
        output_filename = f"edited_{filename}"
        output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
        doc.save(output_path)
        doc.close()
        
        return jsonify({
            'message': 'PDF edited successfully',
            'edited_filename': output_filename
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/convert/pdf-to-word', methods=['POST'])
def pdf_to_word():
    """Convert PDF to Word document"""
    data = request.json
    filename = data.get('filename')
    
    if not filename:
        return jsonify({'error': 'Missing filename'}), 400
    
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(input_path):
        return jsonify({'error': 'File not found'}), 404
    
    try:
        # Create Word document
        doc = Document()
        
        # Open PDF and extract text
        pdf_doc = fitz.open(input_path)
        
        for page_num in range(len(pdf_doc)):
            page = pdf_doc[page_num]
            text = page.get_text()
            
            # Add page text to Word document
            if text.strip():
                doc.add_paragraph(text)
            
            # Add page break if not the last page
            if page_num < len(pdf_doc) - 1:
                doc.add_page_break()
        
        pdf_doc.close()
        
        # Save Word document
        output_filename = f"{filename.rsplit('.', 1)[0]}.docx"
        output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
        doc.save(output_path)
        
        return jsonify({
            'message': 'PDF converted to Word successfully',
            'word_filename': output_filename
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/convert/word-to-pdf', methods=['POST'])
def word_to_pdf():
    """Convert Word document to PDF"""
    data = request.json
    filename = data.get('filename')
    
    if not filename:
        return jsonify({'error': 'Missing filename'}), 400
    
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(input_path):
        return jsonify({'error': 'File not found'}), 404
    
    try:
        # Read Word document
        doc = Document(input_path)
        
        # Create PDF using reportlab
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
        from reportlab.lib.styles import getSampleStyleSheet
        
        output_filename = f"{filename.rsplit('.', 1)[0]}.pdf"
        output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
        
        pdf_doc = SimpleDocTemplate(output_path, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Extract text from Word document
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                p = Paragraph(paragraph.text, styles['Normal'])
                story.append(p)
                story.append(Spacer(1, 12))
        
        pdf_doc.build(story)
        
        return jsonify({
            'message': 'Word document converted to PDF successfully',
            'pdf_filename': output_filename
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/split', methods=['POST'])
def split_pdf():
    """Split PDF into multiple files"""
    data = request.json
    filename = data.get('filename')
    ranges = data.get('ranges')  # e.g., ["1-2", "3", "4-6"]
    
    if not filename or not ranges:
        return jsonify({'error': 'Missing filename or ranges'}), 400
    
    input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if not os.path.exists(input_path):
        return jsonify({'error': 'File not found'}), 404
    
    try:
        split_files = []
        reader = PdfReader(input_path)
        total_pages = len(reader.pages)
        
        for idx, page_range in enumerate(ranges):
            writer = PdfWriter()
            
            # Parse page range
            if '-' in page_range:
                start, end = map(int, page_range.split('-'))
                start = max(1, start)
                end = min(total_pages, end)
            else:
                start = end = int(page_range)
                if start < 1 or start > total_pages:
                    continue
            
            # Add pages to writer
            for page_num in range(start - 1, end):
                writer.add_page(reader.pages[page_num])
            
            # Save split file
            output_filename = f"split_{idx + 1}_{filename}"
            output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
            
            with open(output_path, 'wb') as output_file:
                writer.write(output_file)
            
            split_files.append(output_filename)
        
        return jsonify({
            'message': 'PDF split successfully',
            'split_files': split_files
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/merge', methods=['POST'])
def merge_pdfs():
    """Merge multiple PDFs into one"""
    data = request.json
    filenames = data.get('filenames', [])
    
    if not filenames or len(filenames) < 2:
        return jsonify({'error': 'Need at least 2 files to merge'}), 400
    
    try:
        writer = PdfWriter()
        
        for filename in filenames:
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            if not os.path.exists(input_path):
                continue
            
            reader = PdfReader(input_path)
            for page in reader.pages:
                writer.add_page(page)
        
        # Save merged PDF
        output_filename = f"merged_{uuid.uuid4()}.pdf"
        output_path = os.path.join(app.config['PROCESSED_FOLDER'], output_filename)
        
        with open(output_path, 'wb') as output_file:
            writer.write(output_file)
        
        return jsonify({
            'message': 'PDFs merged successfully',
            'merged_filename': output_filename
        }), 200
    
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/delete/<filename>', methods=['DELETE'])
def delete_file(filename):
    """Delete a file"""
    file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
    if os.path.exists(file_path):
        os.remove(file_path)
        return jsonify({'message': 'File deleted successfully'}), 200
    return jsonify({'error': 'File not found'}), 404

@app.route('/')
def index():
    return jsonify({'message': 'PDF Editor Flask backend running'})

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True, threaded=True) 